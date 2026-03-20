"""
Export HTML content to proper DOCX with native Word math equations (OMML).
Uses python-docx + latex2mathml + lxml XSLT for LaTeX -> MathML -> OMML pipeline.
"""
import re
import io
from html.parser import HTMLParser
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from lxml import etree
import latex2mathml.converter

# Load XSLT once at module level
_xslt = etree.parse("/app/backend/MML2OMML.XSL")
_transform = etree.XSLT(_xslt)

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _fix_matrix_delimiters(omml_root):
    """
    Fix matrix delimiters in OMML: replace plain text '(' + matrix + ')' 
    with proper <m:d> delimiter that auto-scales.
    """
    children = list(omml_root)
    i = 0
    while i < len(children) - 2:
        # Look for pattern: <m:r>(</m:r> <m:m>...</m:m> <m:r>)</m:r>
        c1, c2, c3 = children[i], children[i + 1], children[i + 2]

        tag1 = c1.tag.split('}')[-1] if '}' in c1.tag else c1.tag
        tag2 = c2.tag.split('}')[-1] if '}' in c2.tag else c2.tag
        tag3 = c3.tag.split('}')[-1] if '}' in c3.tag else c3.tag

        if tag1 == 'r' and tag2 == 'm' and tag3 == 'r':
            # Check if c1 is "(" and c3 is ")"
            t1 = c1.find(f'{{{OMML_NS}}}t')
            t3 = c3.find(f'{{{OMML_NS}}}t')
            t1_text = (t1.text or '').strip() if t1 is not None else ''
            t3_text = (t3.text or '').strip() if t3 is not None else ''

            if t1_text in ('(', '[', '|') and t3_text in (')', ']', '|'):
                # Create proper <m:d> delimiter
                d = etree.SubElement(omml_root, f'{{{OMML_NS}}}d')
                dPr = etree.SubElement(d, f'{{{OMML_NS}}}dPr')
                begChr = etree.SubElement(dPr, f'{{{OMML_NS}}}begChr')
                begChr.set(f'{{{OMML_NS}}}val', t1_text)
                endChr = etree.SubElement(dPr, f'{{{OMML_NS}}}endChr')
                endChr.set(f'{{{OMML_NS}}}val', t3_text)
                e = etree.SubElement(d, f'{{{OMML_NS}}}e')
                e.append(c2)  # Move matrix inside delimiter

                # Remove old elements and insert delimiter
                omml_root.remove(c1)
                omml_root.remove(c3)
                # Move d to the position where c1 was
                children = list(omml_root)
                idx = children.index(d)
                if idx != i:
                    omml_root.remove(d)
                    omml_root.insert(i, d)

                children = list(omml_root)
                continue

        i += 1

    return omml_root


def latex_to_omml(latex_str):
    """Convert LaTeX string to OMML etree element with fixed delimiters"""
    try:
        mathml = latex2mathml.converter.convert(latex_str)
        mathml_tree = etree.fromstring(mathml.encode('utf-8'))
        omml_tree = _transform(mathml_tree)
        omml_root = omml_tree.getroot()
        # Fix matrix delimiters
        _fix_matrix_delimiters(omml_root)
        return omml_root
    except Exception:
        return None


def _add_shading(paragraph, color_hex="1E3A5F"):
    """Add background shading to a paragraph"""
    pPr = paragraph._element.get_or_add_pPr()
    shd = etree.SubElement(pPr, qn('w:shd'))
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)


def _set_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


class HTMLToDocxParser(HTMLParser):
    """Parse HTML and build a python-docx Document with proper math rendering"""

    def __init__(self, doc):
        super().__init__()
        self.doc = doc
        self.current_paragraph = None
        self.in_table = False
        self.table_data = []
        self.current_row = []
        self.current_cell_text = ""
        self.in_th = False
        self.in_td = False
        self.tag_stack = []
        self.skip_content = False
        self.pending_text = ""
        self.in_strong = False
        self.in_em = False
        self.in_indent = False
        self.last_was_heading = False

    def _flush_text(self):
        """Flush pending text to current paragraph, handling LaTeX math"""
        if not self.pending_text or not self.current_paragraph:
            return

        text = self.pending_text
        self.pending_text = ""

        # Split text by $...$ LaTeX patterns
        parts = re.split(r'(\$\$[\s\S]+?\$\$|\$[^\$\n]+?\$)', text)
        for part in parts:
            if not part:
                continue

            # Display math $$...$$
            if part.startswith('$$') and part.endswith('$$'):
                latex = part[2:-2].strip()
                omml = latex_to_omml(latex)
                if omml is not None:
                    self.current_paragraph._element.append(omml)
                else:
                    run = self.current_paragraph.add_run(latex)
                    run.italic = True

            # Inline math $...$
            elif part.startswith('$') and part.endswith('$') and len(part) > 2:
                latex = part[1:-1].strip()
                omml = latex_to_omml(latex)
                if omml is not None:
                    self.current_paragraph._element.append(omml)
                else:
                    run = self.current_paragraph.add_run(latex)
                    run.italic = True

            # Regular text
            else:
                cleaned = part.replace('\n', ' ').strip()
                if cleaned:
                    run = self.current_paragraph.add_run(cleaned)
                    run.bold = self.in_strong
                    run.italic = self.in_em
                    run.font.size = Pt(11)
                    run.font.name = 'Cambria'

    def _new_paragraph(self, is_heading=False):
        """Create a new paragraph, reusing current if empty"""
        if self.current_paragraph and not self.current_paragraph.text.strip():
            # Check if it has any runs or math elements
            has_content = len(self.current_paragraph.runs) > 0
            omml_in_p = self.current_paragraph._element.findall(f'.//{{{OMML_NS}}}oMath')
            if not has_content and not omml_in_p:
                return self.current_paragraph

        p = self.doc.add_paragraph()
        _set_spacing(p, before=1, after=1, line=1.15)
        return p

    def handle_starttag(self, tag, attrs):
        attrs_dict = dict(attrs)
        self.tag_stack.append(tag)

        if tag in ('svg', 'script', 'style'):
            self.skip_content = True
            return

        if self.in_table and tag in ('td', 'th'):
            self.in_th = tag == 'th'
            self.in_td = tag == 'td'
            self.current_cell_text = ""
            return

        if self.in_table and tag == 'tr':
            self.current_row = []
            return

        if tag == 'table':
            self._flush_text()
            self.in_table = True
            self.table_data = []
            return

        if self.in_table:
            return

        style = attrs_dict.get('style', '')

        if tag == 'h1':
            self._flush_text()
            self.current_paragraph = self.doc.add_heading(level=1)
            _set_spacing(self.current_paragraph, before=6, after=3)
            self.last_was_heading = True

        elif tag == 'h2':
            self._flush_text()
            self.current_paragraph = self.doc.add_paragraph()
            _set_spacing(self.current_paragraph, before=8, after=4)
            # Add blue background
            _add_shading(self.current_paragraph, "1E3A5F")
            if 'text-align:center' in style or 'text-align: center' in style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self.last_was_heading = True

        elif tag == 'h3':
            self._flush_text()
            self.current_paragraph = self.doc.add_heading(level=3)
            _set_spacing(self.current_paragraph, before=6, after=2)
            self.last_was_heading = True

        elif tag == 'p':
            self._flush_text()
            self.current_paragraph = self._new_paragraph()
            if 'text-align:center' in style or 'text-align: center' in style:
                self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif tag == 'div':
            if 'margin-left' in style:
                self._flush_text()
                self.current_paragraph = self._new_paragraph()
                pf = self.current_paragraph.paragraph_format
                pf.left_indent = Cm(1.2)
                self.in_indent = True
            elif 'margin-bottom' in style:
                # Question div — just ensure we have a paragraph
                self._flush_text()

        elif tag == 'br':
            if self.current_paragraph:
                self._flush_text()
                self.current_paragraph.add_run('\n')

        elif tag == 'hr':
            self._flush_text()
            hr_p = self.doc.add_paragraph()
            _set_spacing(hr_p, before=4, after=4)
            hr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hr_p.add_run('─' * 60)
            run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
            run.font.size = Pt(8)
            self.current_paragraph = None

        elif tag == 'img':
            self._flush_text()
            src = attrs_dict.get('src', '')
            if src.startswith('data:image/'):
                try:
                    import base64
                    header, b64data = src.split(',', 1)
                    img_bytes = base64.b64decode(b64data)
                    img_stream = io.BytesIO(img_bytes)
                    if not self.current_paragraph:
                        self.current_paragraph = self._new_paragraph()
                    self.current_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = self.current_paragraph.add_run()
                    run.add_picture(img_stream, width=Cm(8))
                    self.current_paragraph = self._new_paragraph()
                except Exception:
                    pass

        elif tag in ('strong', 'b'):
            self._flush_text()
            self.in_strong = True

        elif tag in ('em', 'i'):
            self._flush_text()
            self.in_em = True

    def handle_endtag(self, tag):
        if self.tag_stack and self.tag_stack[-1] == tag:
            self.tag_stack.pop()

        if tag in ('svg', 'script', 'style'):
            self.skip_content = False
            return

        if tag in ('strong', 'b'):
            self._flush_text()
            self.in_strong = False
            return

        if tag in ('em', 'i'):
            self._flush_text()
            self.in_em = False
            return

        if self.in_table:
            if tag in ('td', 'th'):
                self.current_row.append({
                    'text': self.current_cell_text.strip(),
                    'is_header': self.in_th,
                })
                self.in_th = False
                self.in_td = False

            elif tag == 'tr':
                if self.current_row:
                    self.table_data.append(self.current_row)
                self.current_row = []

            elif tag == 'table':
                self.in_table = False
                self._build_table()

            return

        if tag == 'div' and self.in_indent:
            self._flush_text()
            self.in_indent = False
            self.current_paragraph = None

        elif tag in ('p',):
            self._flush_text()
            self.current_paragraph = None

        elif tag in ('h1', 'h2', 'h3'):
            self._flush_text()
            self.current_paragraph = None

    def handle_data(self, data):
        if self.skip_content:
            return

        if self.in_table and (self.in_th or self.in_td):
            self.current_cell_text += data
            return

        text = data
        if text.strip():
            if self.current_paragraph is None:
                self.current_paragraph = self._new_paragraph()

            # For h2 headings (blue bg), style the text as white bold
            pPr = self.current_paragraph._element.find(f'{{{W_NS}}}pPr')
            is_shaded = False
            if pPr is not None:
                shd = pPr.find(f'{{{W_NS}}}shd')
                is_shaded = shd is not None

            if is_shaded and not self.pending_text:
                # This is heading text on blue background
                self.pending_text += text
            else:
                self.pending_text += text

    def _finalize_shaded_paragraphs(self):
        """Style text in shaded (blue bg) paragraphs as white bold"""
        for p in self.doc.paragraphs:
            pPr = p._element.find(f'{{{W_NS}}}pPr')
            if pPr is not None:
                shd = pPr.find(f'{{{W_NS}}}shd')
                if shd is not None and shd.get(qn('w:fill')) == '1E3A5F':
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True
                        run.font.size = Pt(13)
                        run.font.name = 'Cambria'

    def _build_table(self):
        """Build a Word table from parsed table data"""
        if not self.table_data:
            return

        max_cols = max(len(row) for row in self.table_data)
        if max_cols == 0:
            return

        table = self.doc.add_table(rows=len(self.table_data), cols=max_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        for i, row_data in enumerate(self.table_data):
            for j, cell_data in enumerate(row_data):
                if j >= max_cols:
                    break
                cell = table.cell(i, j)
                # Clear default paragraph
                cell.paragraphs[0].text = ''
                run = cell.paragraphs[0].add_run(cell_data['text'])
                run.font.size = Pt(10)
                run.font.name = 'Cambria'
                _set_spacing(cell.paragraphs[0], before=1, after=1)

                if cell_data['is_header'] or i == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    bg = etree.SubElement(tc_pr, qn('w:shd'))
                    bg.set(qn('w:val'), 'clear')
                    bg.set(qn('w:color'), 'auto')
                    bg.set(qn('w:fill'), '1E3A5F')

        self.table_data = []
        self.current_paragraph = None


def _remove_empty_paragraphs(doc):
    """Remove consecutive empty paragraphs, keep at most one"""
    body = doc.element.body
    paragraphs = body.findall(f'{{{W_NS}}}p')
    prev_empty = False
    for p in paragraphs:
        # Check if paragraph has any text or math content
        text = ''.join(t.text or '' for t in p.iter(f'{{{W_NS}}}t'))
        has_math = p.find(f'.//{{{OMML_NS}}}oMath') is not None

        if not text.strip() and not has_math:
            if prev_empty:
                body.remove(p)
            else:
                prev_empty = True
        else:
            prev_empty = False


def html_to_docx(html_content, title="Dokumen"):
    """Convert HTML content with LaTeX math to a proper DOCX file"""
    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    pf.line_spacing = 1.15

    # Heading styles
    for level in (1, 2, 3):
        h_style = doc.styles[f'Heading {level}']
        h_style.font.name = 'Cambria'
        h_style.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Clean HTML
    html_content = re.sub(r'\[DIAGRAM:[^\]]+\]', '[Lihat Diagram di Aplikasi]', html_content)

    # Parse and build document
    parser = HTMLToDocxParser(doc)
    parser.feed(html_content)
    parser._flush_text()

    # Post-process: style shaded paragraphs, remove empties
    parser._finalize_shaded_paragraphs()
    _remove_empty_paragraphs(doc)

    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
